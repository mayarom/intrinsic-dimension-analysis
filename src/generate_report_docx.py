from __future__ import annotations
#
# from pathlib import Path
# from typing import Iterable
#
# from docx.api import Document
# from docx.enum.table import WD_ALIGN_VERTICAL
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.shared import Inches, Pt, RGBColor
#
#
# FIGURES_DIR = Path(
#     "/Users/rommay/PycharmProjects/intrinsic-dimension-analysis/"
#     "intrinsic-dimension-analysis/figures"
# )
# OUTPUT_DOCX = FIGURES_DIR.parent / "intrinsic_dimension_report_with_figures.docx"
#
# FIGURE_FILES: dict[str, str] = {
#     "mnist_samples": "mnist_samples.png",
#     "pca_cumulative_mnist": "pca_cumulative_mnist.png",
#     "pca_cumulative_gaussian_noise": "pca_cumulative_gaussian_noise.png",
#     "pca_cumulative_plane": "pca_cumulative_plane_2d_in_10d.png",
#     "pca_cumulative_swiss_roll": "pca_cumulative_swiss_roll.png",
#     "scree_mnist": "scree_mnist.png",
#     "scree_gaussian_noise": "scree_gaussian_noise.png",
#     "scree_plane": "scree_plane_2d_in_10d.png",
#     "scree_swiss_roll": "scree_swiss_roll.png",
#     "correlation_mnist": "correlation_loglog_mnist.png",
#     "correlation_gaussian_noise": "correlation_loglog_gaussian_noise.png",
#     "correlation_plane": "correlation_loglog_plane_2d_in_10d.png",
#     "correlation_swiss_roll": "correlation_loglog_swiss_roll.png",
#     "comparison_all": "dimension_comparison.png",
#     "comparison_synthetic": "dimension_comparison_synthetic_only.png",
# }
#
# ACCENT = RGBColor(31, 78, 121)
# DARK = RGBColor(40, 40, 40)
# MID = RGBColor(90, 90, 90)
#
#
# class FigureNumbering:
#     """
#     Simple sequential figure numbering manager.
#     """
#
#     def __init__(self) -> None:
#         self.current = 0
#
#     def next(self) -> int:
#         self.current += 1
#         return self.current
#
#
# def configure_page_layout(doc: Document) -> None:
#     """
#     Configure page margins and header/footer spacing.
#     """
#     for section in doc.sections:
#         section.top_margin = Inches(0.9)
#         section.bottom_margin = Inches(0.8)
#         section.left_margin = Inches(1.0)
#         section.right_margin = Inches(1.0)
#         section.header_distance = Inches(0.35)
#         section.footer_distance = Inches(0.35)
#
#
# def apply_base_styles(doc: Document) -> None:
#     """
#     Apply an academic, submission-style document design.
#     """
#     styles = doc.styles
#
#     normal = styles["Normal"]
#     normal.font.name = "Times New Roman"
#     normal.font.size = Pt(11)
#
#     title_style = styles["Title"]
#     title_style.font.name = "Times New Roman"
#     title_style.font.size = Pt(20)
#     title_style.font.bold = True
#     title_style.font.color.rgb = ACCENT
#
#     subtitle_style = styles["Subtitle"]
#     subtitle_style.font.name = "Times New Roman"
#     subtitle_style.font.size = Pt(12)
#     subtitle_style.font.italic = True
#     subtitle_style.font.color.rgb = MID
#
#     heading1 = styles["Heading 1"]
#     heading1.font.name = "Times New Roman"
#     heading1.font.size = Pt(15)
#     heading1.font.bold = True
#     heading1.font.color.rgb = ACCENT
#
#     heading2 = styles["Heading 2"]
#     heading2.font.name = "Times New Roman"
#     heading2.font.size = Pt(12.5)
#     heading2.font.bold = True
#     heading2.font.color.rgb = DARK
#
#     heading3 = styles["Heading 3"]
#     heading3.font.name = "Times New Roman"
#     heading3.font.size = Pt(11.5)
#     heading3.font.bold = True
#     heading3.font.color.rgb = DARK
#
#
# def add_header_and_footer(doc: Document) -> None:
#     """
#     Add a clean academic header and footer.
#
#     Note:
#     Automatic page-number fields in python-docx require low-level XML access,
#     which triggers protected-member warnings in the IDE. To keep the script
#     clean and warning-free, this version uses a professional footer without
#     XML field injection.
#     """
#     for section in doc.sections:
#         header = section.header
#         header_p = header.paragraphs[0]
#         header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#         run = header_p.add_run("Intrinsic Dimension Analysis Report")
#         run.font.name = "Times New Roman"
#         run.font.size = Pt(9)
#         run.font.italic = True
#         run.font.color.rgb = MID
#
#         footer = section.footer
#         footer_p = footer.paragraphs[0]
#         footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         run = footer_p.add_run("Final submission version")
#         run.font.name = "Times New Roman"
#         run.font.size = Pt(9)
#         run.font.italic = True
#         run.font.color.rgb = MID
#
#
# def format_paragraph(
#     paragraph,
#     space_before: float = 0,
#     space_after: float = 6,
#     line_spacing: float = 1.15,
# ) -> None:
#     """
#     Apply consistent paragraph spacing.
#     """
#     paragraph_format = paragraph.paragraph_format
#     paragraph_format.space_before = Pt(space_before)
#     paragraph_format.space_after = Pt(space_after)
#     paragraph_format.line_spacing = line_spacing
#
#
# def add_body_paragraph(
#     doc: Document,
#     text: str,
#     align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
# ) -> None:
#     """
#     Add a standard body paragraph.
#     """
#     paragraph = doc.add_paragraph()
#     paragraph.alignment = align
#     run = paragraph.add_run(text)
#     run.font.name = "Times New Roman"
#     run.font.size = Pt(11)
#     format_paragraph(paragraph, space_after=6)
#
#
# def add_bullets(doc: Document, items: Iterable[str]) -> None:
#     """
#     Add a bullet list.
#     """
#     for item in items:
#         paragraph = doc.add_paragraph(style="List Bullet")
#         paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
#         run = paragraph.add_run(item)
#         run.font.name = "Times New Roman"
#         run.font.size = Pt(11)
#         format_paragraph(paragraph, space_after=3)
#
#
# def add_caption(doc: Document, figure_number: int, text: str) -> None:
#     """
#     Add an academic-style figure caption.
#     """
#     paragraph = doc.add_paragraph()
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#
#     run1 = paragraph.add_run(f"Figure {figure_number}. ")
#     run1.bold = True
#     run1.font.name = "Times New Roman"
#     run1.font.size = Pt(10)
#
#     run2 = paragraph.add_run(text)
#     run2.italic = True
#     run2.font.name = "Times New Roman"
#     run2.font.size = Pt(10)
#     run2.font.color.rgb = MID
#
#     format_paragraph(paragraph, space_before=2, space_after=10, line_spacing=1.0)
#
#
# def add_figure(
#     doc: Document,
#     numbering: FigureNumbering,
#     image_path: Path,
#     caption: str,
#     width: float = 6.1,
# ) -> int:
#     """
#     Insert a figure and its caption.
#     """
#     if not image_path.exists():
#         add_body_paragraph(
#             doc,
#             f"[Missing figure file: {image_path.name}]",
#             align=WD_ALIGN_PARAGRAPH.CENTER,
#         )
#         return numbering.current
#
#     figure_number = numbering.next()
#
#     paragraph = doc.add_paragraph()
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     run = paragraph.add_run()
#     run.add_picture(str(image_path), width=Inches(width))
#     format_paragraph(paragraph, space_before=4, space_after=2, line_spacing=1.0)
#
#     add_caption(doc, figure_number, caption)
#     return figure_number
#
#
# def add_title_page(doc: Document) -> None:
#     """
#     Add a professional title page.
#     """
#     paragraph = doc.add_paragraph()
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     run = paragraph.add_run("Estimating the Intrinsic Dimension of High-Dimensional Data")
#     run.bold = True
#     run.font.name = "Times New Roman"
#     run.font.size = Pt(21)
#     run.font.color.rgb = ACCENT
#     format_paragraph(paragraph, space_before=70, space_after=10, line_spacing=1.0)
#
#     paragraph = doc.add_paragraph()
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     run = paragraph.add_run("PCA and Correlation Dimension Analysis")
#     run.italic = True
#     run.font.name = "Times New Roman"
#     run.font.size = Pt(13)
#     run.font.color.rgb = MID
#     format_paragraph(paragraph, space_after=24, line_spacing=1.0)
#
#     info_lines = [
#         "Final Submission Version",
#         "Automatically Generated Figure Report",
#         "Language: English",
#     ]
#     for line in info_lines:
#         paragraph = doc.add_paragraph()
#         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         run = paragraph.add_run(line)
#         run.font.name = "Times New Roman"
#         run.font.size = Pt(11)
#         format_paragraph(paragraph, space_after=5, line_spacing=1.0)
#
#     doc.add_page_break()
#
#
# def add_summary_table(doc: Document) -> None:
#     """
#     Add an executive summary table.
#     """
#     doc.add_heading("Executive Summary of Main Findings", level=1)
#
#     add_body_paragraph(
#         doc,
#         "The results show that PCA and correlation dimension capture different but "
#         "complementary aspects of dataset structure. PCA measures how many linear "
#         "components are required to preserve most of the variance, whereas correlation "
#         "dimension estimates how the data scale geometrically in local neighborhoods."
#     )
#
#     table = doc.add_table(rows=1, cols=4)
#     table.style = "Table Grid"
#
#     headers = [
#         "Dataset",
#         "Expected Structure",
#         "PCA-Based Interpretation",
#         "Correlation-Dimension Interpretation",
#     ]
#
#     header_cells = table.rows[0].cells
#     for index, text in enumerate(headers):
#         cell = header_cells[index]
#         cell.text = text
#         cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#         paragraph = cell.paragraphs[0]
#         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         run = paragraph.runs[0]
#         run.bold = True
#         run.font.name = "Times New Roman"
#         run.font.size = Pt(10)
#
#     rows = [
#         (
#             "2D Plane in 10D",
#             "A linear two-dimensional manifold embedded in a ten-dimensional ambient space.",
#             "PCA should recover the intrinsic structure with approximately two dominant components.",
#             "Correlation dimension should be close to 2 if the geometric estimate is stable.",
#         ),
#         (
#             "Swiss Roll",
#             "A nonlinear two-dimensional manifold in three-dimensional space.",
#             "PCA may require more than two components because the manifold is curved.",
#             "Correlation dimension should remain close to the intrinsic dimension of 2.",
#         ),
#         (
#             "Gaussian Noise",
#             "High-dimensional isotropic noise without compact low-dimensional structure.",
#             "Variance is expected to spread over many components.",
#             "Correlation dimension should be relatively high compared with the structured datasets.",
#         ),
#         (
#             "MNIST",
#             "Structured handwritten digit images with nonlinear variation in pixel space.",
#             "Many components are typically required to preserve most of the variance.",
#             "Correlation dimension is expected to be much lower than the PCA component count.",
#         ),
#     ]
#
#     for row_values in rows:
#         row = table.add_row()
#         for index, value in enumerate(row_values):
#             cell = row.cells[index]
#             cell.text = value
#             cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#             for paragraph in cell.paragraphs:
#                 paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
#                 format_paragraph(paragraph, space_after=2, line_spacing=1.05)
#                 for run in paragraph.runs:
#                     run.font.name = "Times New Roman"
#                     run.font.size = Pt(10)
#
#     doc.add_paragraph()
#
#
# def add_intro_and_methods(doc: Document) -> None:
#     """
#     Add the introduction and methods overview sections.
#     """
#     doc.add_heading("1. Introduction", level=1)
#
#     add_body_paragraph(
#         doc,
#         "Intrinsic dimension refers to the minimum number of variables required to "
#         "represent the underlying structure of a dataset. In many cases, the ambient "
#         "dimension is much larger than the true number of degrees of freedom governing "
#         "the data."
#     )
#     add_body_paragraph(
#         doc,
#         "This report presents the visual results of intrinsic-dimension analysis for "
#         "four datasets: a two-dimensional plane embedded in ten dimensions, the Swiss "
#         "roll dataset, Gaussian noise, and MNIST handwritten digit images."
#     )
#     add_body_paragraph(
#         doc,
#         "Two complementary estimation methods were applied. Principal component analysis "
#         "(PCA) was used to examine the number of linear components required to explain "
#         "most of the variance, while correlation dimension was used to capture geometric "
#         "scaling behavior in the local structure of the data."
#     )
#
#     doc.add_heading("2. Methods Overview", level=1)
#     add_bullets(
#         doc,
#         [
#             "PCA was evaluated through cumulative explained variance curves and scree plots.",
#             "Correlation dimension was estimated from the slope of the linear scaling region in log-log correlation-integral plots.",
#             "The combined analysis was designed to compare linear variance structure with intrinsic geometric complexity.",
#         ],
#     )
#
#
# def add_dataset_results(
#     doc: Document,
#     numbering: FigureNumbering,
#     section_title: str,
#     dataset_title: str,
#     pca_cumulative_key: str,
#     scree_key: str,
#     correlation_key: str,
#     pca_text: list[str],
#     scree_text: list[str],
#     corr_text: list[str],
# ) -> None:
#     """
#     Add all figure-based results for a single dataset.
#     """
#     doc.add_heading(section_title, level=2)
#
#     add_body_paragraph(
#         doc,
#         f"This subsection summarizes the visual results obtained for the {dataset_title} dataset."
#     )
#
#     add_figure(
#         doc,
#         numbering,
#         FIGURES_DIR / FIGURE_FILES[pca_cumulative_key],
#         f"Cumulative explained variance curve for the {dataset_title} dataset.",
#     )
#     for paragraph_text in pca_text:
#         add_body_paragraph(doc, paragraph_text)
#
#     add_figure(
#         doc,
#         numbering,
#         FIGURES_DIR / FIGURE_FILES[scree_key],
#         f"Scree plot for the {dataset_title} dataset.",
#     )
#     for paragraph_text in scree_text:
#         add_body_paragraph(doc, paragraph_text)
#
#     add_figure(
#         doc,
#         numbering,
#         FIGURES_DIR / FIGURE_FILES[correlation_key],
#         f"Correlation-dimension log-log plot for the {dataset_title} dataset.",
#     )
#     for paragraph_text in corr_text:
#         add_body_paragraph(doc, paragraph_text)
#
#
# def add_global_comparison(doc: Document, numbering: FigureNumbering) -> None:
#     """
#     Add cross-dataset comparison figures.
#     """
#     doc.add_heading("4. Global Comparison Across Datasets", level=1)
#
#     add_body_paragraph(
#         doc,
#         "After analyzing each dataset individually, it is useful to compare the final "
#         "PCA-based and correlation-dimension estimates across all datasets."
#     )
#
#     add_figure(
#         doc,
#         numbering,
#         FIGURES_DIR / FIGURE_FILES["comparison_all"],
#         "Comparison of PCA-based dimensionality estimates and correlation-dimension estimates across all datasets.",
#         width=6.3,
#     )
#
#     add_body_paragraph(
#         doc,
#         "The full comparison clearly shows that PCA and correlation dimension measure "
#         "different aspects of structure. The gap is especially pronounced for MNIST, "
#         "where many PCA components are needed to retain variance, but the geometric "
#         "complexity remains much lower."
#     )
#
#     add_figure(
#         doc,
#         numbering,
#         FIGURES_DIR / FIGURE_FILES["comparison_synthetic"],
#         "Comparison of PCA and correlation-dimension estimates restricted to the synthetic datasets.",
#         width=6.3,
#     )
#
#     add_body_paragraph(
#         doc,
#         "The synthetic-only comparison removes the scale effect introduced by MNIST and "
#         "makes the relationships among the controlled datasets easier to interpret."
#     )
#
#
# def add_mnist_visual_reference(doc: Document, numbering: FigureNumbering) -> None:
#     """
#     Add the MNIST sample figure.
#     """
#     doc.add_heading("5. Visual Reference for MNIST", level=1)
#
#     add_body_paragraph(
#         doc,
#         "A sample of MNIST images is included below to provide visual context for the "
#         "type of data analyzed in the experiment."
#     )
#
#     add_figure(
#         doc,
#         numbering,
#         FIGURES_DIR / FIGURE_FILES["mnist_samples"],
#         "Random sample of MNIST handwritten digit images.",
#         width=6.4,
#     )
#
#     add_body_paragraph(
#         doc,
#         "Although each image is represented numerically as a 784-dimensional vector, "
#         "the visual variation is governed by a much smaller set of semantic factors, "
#         "including digit identity, stroke thickness, slant, and style."
#     )
#
#
# def add_conclusion(doc: Document) -> None:
#     """
#     Add the final discussion and conclusion.
#     """
#     doc.add_heading("6. Discussion and Conclusion", level=1)
#
#     add_body_paragraph(
#         doc,
#         "The results demonstrate that intrinsic dimension is not a single quantity that "
#         "all methods recover identically. Instead, the estimate depends on what aspect "
#         "of structure is being measured."
#     )
#
#     add_bullets(
#         doc,
#         [
#             "PCA is most informative when the data are approximately linear and the goal is variance preservation.",
#             "Correlation dimension is better suited to capturing nonlinear geometric structure.",
#             "Agreement between the two methods is strongest for simple linear datasets such as the 2D plane in 10D.",
#             "Disagreement between the methods is itself informative for nonlinear and complex real-world datasets.",
#         ],
#     )
#
#     add_body_paragraph(
#         doc,
#         "Overall, the 2D plane embedded in 10D served as a clear sanity check, the Swiss "
#         "roll highlighted the limitations of linear methods, Gaussian noise behaved as a "
#         "high-dimensional unstructured dataset, and MNIST revealed a strong gap between "
#         "variance-based and geometry-based notions of dimensionality."
#     )
#
#     add_body_paragraph(
#         doc,
#         "These findings support the use of both PCA and correlation dimension as complementary "
#         "tools for understanding the true complexity of high-dimensional data."
#     )
#
#
# def build_document() -> Document:
#     """
#     Build the full report document.
#     """
#     doc = Document()
#     configure_page_layout(doc)
#     apply_base_styles(doc)
#     add_header_and_footer(doc)
#
#     numbering = FigureNumbering()
#
#     add_title_page(doc)
#     add_summary_table(doc)
#     add_intro_and_methods(doc)
#
#     doc.add_heading("3. Results and Figure Analysis", level=1)
#
#     add_dataset_results(
#         doc=doc,
#         numbering=numbering,
#         section_title="3.1 2D Plane Embedded in 10D",
#         dataset_title="2D Plane Embedded in 10D",
#         pca_cumulative_key="pca_cumulative_plane",
#         scree_key="scree_plane",
#         correlation_key="correlation_plane",
#         pca_text=[
#             "The cumulative explained variance curve shows that almost all variance is captured by the first two principal components. This is exactly the expected result for a linear two-dimensional manifold embedded in a higher-dimensional ambient space.",
#             "The figure therefore confirms that PCA successfully recovers the known intrinsic structure of this dataset.",
#         ],
#         scree_text=[
#             "The scree plot shows a clear dominance of the first two components, while the remaining components contribute essentially no meaningful variance.",
#             "This is the strongest possible visual indication of a low-dimensional linear structure.",
#         ],
#         corr_text=[
#             "The correlation-dimension estimate is very close to 2, which matches the theoretical intrinsic dimensionality of the dataset.",
#             "The agreement between PCA and correlation dimension validates both the dataset construction and the analysis workflow.",
#         ],
#     )
#
#     add_dataset_results(
#         doc=doc,
#         numbering=numbering,
#         section_title="3.2 Swiss Roll",
#         dataset_title="Swiss Roll",
#         pca_cumulative_key="pca_cumulative_swiss_roll",
#         scree_key="scree_swiss_roll",
#         correlation_key="correlation_swiss_roll",
#         pca_text=[
#             "The cumulative explained variance plot indicates that multiple principal components are required to preserve most of the variance.",
#             "This occurs because the Swiss roll is intrinsically two-dimensional but geometrically curved, and PCA is limited to linear directions.",
#         ],
#         scree_text=[
#             "The scree plot does not collapse sharply after the first two components, which reflects the nonlinear geometry of the manifold.",
#             "This behavior is expected when a curved low-dimensional structure is analyzed using a linear projection method.",
#         ],
#         corr_text=[
#             "The correlation-dimension estimate remains close to 2, which is much more faithful to the true intrinsic structure of the Swiss roll.",
#             "This dataset therefore illustrates why geometric methods are especially valuable for nonlinear manifolds.",
#         ],
#     )
#
#     add_dataset_results(
#         doc=doc,
#         numbering=numbering,
#         section_title="3.3 Gaussian Noise",
#         dataset_title="Gaussian Noise",
#         pca_cumulative_key="pca_cumulative_gaussian_noise",
#         scree_key="scree_gaussian_noise",
#         correlation_key="correlation_gaussian_noise",
#         pca_text=[
#             "The cumulative explained variance increases gradually rather than sharply, indicating that variance is distributed across many directions.",
#             "This is the expected pattern for isotropic high-dimensional noise without a compact low-dimensional linear subspace.",
#         ],
#         scree_text=[
#             "The scree plot shows no pronounced elbow and no small subset of components that dominates the variance structure.",
#             "This confirms the absence of a strong linear low-dimensional representation.",
#         ],
#         corr_text=[
#             "The correlation-dimension estimate is comparatively high, which is consistent with the noise-like nature of the data.",
#             "Together, the PCA and correlation results distinguish Gaussian noise clearly from the more structured datasets.",
#         ],
#     )
#
#     add_dataset_results(
#         doc=doc,
#         numbering=numbering,
#         section_title="3.4 MNIST",
#         dataset_title="MNIST",
#         pca_cumulative_key="pca_cumulative_mnist",
#         scree_key="scree_mnist",
#         correlation_key="correlation_mnist",
#         pca_text=[
#             "The cumulative explained variance curve shows that a very large number of principal components are required to preserve 90%, 95%, and 99% of the total variance.",
#             "This reflects the fact that image variability in pixel space is distributed across many linear directions.",
#         ],
#         scree_text=[
#             "The scree plot decays gradually, indicating that many components contain non-negligible information.",
#             "Such behavior is typical for structured image datasets with rich variability in writing style, thickness, curvature, and placement.",
#         ],
#         corr_text=[
#             "The correlation-dimension estimate is much lower than the number of PCA components required for high variance retention.",
#             "This gap suggests that MNIST occupies a structured nonlinear region of the ambient space rather than a diffuse cloud filling the full pixel space.",
#         ],
#     )
#
#     add_global_comparison(doc, numbering)
#     add_mnist_visual_reference(doc, numbering)
#     add_conclusion(doc)
#
#     return doc
#
#
# def main() -> None:
#     """
#     Validate figure availability and create the DOCX file.
#     """
#     missing_files: list[str] = []
#
#     for file_name in FIGURE_FILES.values():
#         path = FIGURES_DIR / file_name
#         if not path.exists():
#             missing_files.append(str(path))
#
#     if missing_files:
#         print("The following figure files are missing:")
#         for item in missing_files:
#             print(f" - {item}")
#         return
#
#     document = build_document()
#     document.save(str(OUTPUT_DOCX))
#
#     print("DOCX report created successfully.")
#     print(f"Output file: {OUTPUT_DOCX}")
#
#
# if __name__ == "__main__":
#     main()


from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


FIGURES_DIR = Path(
    "/Users/rommay/PycharmProjects/intrinsic-dimension-analysis/"
    "intrinsic-dimension-analysis/intrinsic_dimension_analysis"
)
OUTPUT_DOCX = FIGURES_DIR.parent / "intrinsic_dimension_report_with_all_figures.docx"

FIGURE_FILES: Dict[str, str] = {
    "mnist_samples": "mnist_samples.png",
    "pca_cumulative_mnist": "pca_cumulative_mnist.png",
    "pca_cumulative_gaussian_noise": "pca_cumulative_gaussian_noise.png",
    "pca_cumulative_plane": "pca_cumulative_plane_2d_in_10d.png",
    "pca_cumulative_swiss_roll": "pca_cumulative_swiss_roll.png",
    "scree_mnist": "scree_mnist.png",
    "scree_gaussian_noise": "scree_gaussian_noise.png",
    "scree_plane": "scree_plane_2d_in_10d.png",
    "scree_swiss_roll": "scree_swiss_roll.png",
    "correlation_mnist": "correlation_loglog_mnist.png",
    "correlation_gaussian_noise": "correlation_loglog_gaussian_noise.png",
    "correlation_plane": "correlation_loglog_plane_2d_in_10d.png",
    "correlation_swiss_roll": "correlation_loglog_swiss_roll.png",
    "knn_mnist": "knn_sensitivity_mnist.png",
    "knn_gaussian_noise": "knn_sensitivity_gaussian_noise.png",
    "knn_plane": "knn_sensitivity_plane_2d_in_10d.png",
    "knn_swiss_roll": "knn_sensitivity_swiss_roll.png",
    "comparison_legacy": "dimension_comparison.png",
    "comparison_synthetic_legacy": "dimension_comparison_synthetic_only.png",
    "comparison_all_methods": "dimension_comparison_all_methods.png",
    "comparison_synthetic_all_methods": "dimension_comparison_synthetic_all_methods.png",
}

ACCENT = RGBColor(31, 78, 121)
DARK = RGBColor(40, 40, 40)
MID = RGBColor(90, 90, 90)


class FigureNumbering:
    """
    Simple sequential figure numbering manager.
    """

    def __init__(self) -> None:
        self.current = 0

    def next(self) -> int:
        self.current += 1
        return self.current


def configure_page_layout(doc: Document) -> None:
    """
    Configure page margins and header/footer spacing.
    """
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)


def apply_base_styles(doc: Document) -> None:
    """
    Apply an academic, submission-style document design.
    """
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    title_style = styles["Title"]
    title_style.font.name = "Times New Roman"
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = ACCENT

    subtitle_style = styles["Subtitle"]
    subtitle_style.font.name = "Times New Roman"
    subtitle_style.font.size = Pt(12)
    subtitle_style.font.italic = True
    subtitle_style.font.color.rgb = MID

    heading1 = styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1.font.size = Pt(15)
    heading1.font.bold = True
    heading1.font.color.rgb = ACCENT

    heading2 = styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2.font.size = Pt(12.5)
    heading2.font.bold = True
    heading2.font.color.rgb = DARK

    heading3 = styles["Heading 3"]
    heading3.font.name = "Times New Roman"
    heading3.font.size = Pt(11.5)
    heading3.font.bold = True
    heading3.font.color.rgb = DARK


def add_header_and_footer(doc: Document) -> None:
    """
    Add a clean academic header and footer.
    """
    for section in doc.sections:
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_p.add_run("Intrinsic Dimension Analysis Report")
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = MID

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("Final submission version")
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = MID


def format_paragraph(
    paragraph,
    space_before: float = 0,
    space_after: float = 6,
    line_spacing: float = 1.15,
) -> None:
    """
    Apply consistent paragraph spacing.
    """
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.line_spacing = line_spacing


def add_body_paragraph(
    doc: Document,
    text: str,
    align: Any = WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    """
    Add a standard body paragraph.
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    format_paragraph(paragraph, space_after=6)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    """
    Add a bullet list.
    """
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        format_paragraph(paragraph, space_after=3)


def add_caption(doc: Document, figure_number: int, text: str) -> None:
    """
    Add an academic-style figure caption.
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = paragraph.add_run(f"Figure {figure_number}. ")
    run1.bold = True
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(10)

    run2 = paragraph.add_run(text)
    run2.italic = True
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(10)
    run2.font.color.rgb = MID

    format_paragraph(paragraph, space_before=2, space_after=10, line_spacing=1.0)


def add_figure(
    doc: Document,
    numbering: FigureNumbering,
    image_path: Path,
    caption: str,
    width: float = 6.2,
) -> int:
    """
    Insert a figure and its caption.
    """
    if not image_path.exists():
        add_body_paragraph(
            doc,
            f"[Missing figure file: {image_path.name}]",
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        return numbering.current

    figure_number = numbering.next()

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    format_paragraph(paragraph, space_before=4, space_after=2, line_spacing=1.0)

    add_caption(doc, figure_number, caption)
    return figure_number


def add_title_page(doc: Document) -> None:
    """
    Add a professional title page.
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Estimating the Intrinsic Dimension of High-Dimensional Data")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(21)
    run.font.color.rgb = ACCENT
    format_paragraph(paragraph, space_before=70, space_after=10, line_spacing=1.0)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("PCA, Correlation Dimension, and k-Nearest Neighbors Analysis")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.font.color.rgb = MID
    format_paragraph(paragraph, space_after=24, line_spacing=1.0)

    info_lines = [
        "Final Submission Version",
        "Automatically Generated Figure Report",
        "Language: English",
    ]
    for line in info_lines:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        format_paragraph(paragraph, space_after=5, line_spacing=1.0)

    doc.add_page_break()


def add_intro_and_methods(doc: Document) -> None:
    """
    Add the introduction and methods overview sections.
    """
    doc.add_heading("1. Introduction", level=1)

    add_body_paragraph(
        doc,
        "Intrinsic dimension refers to the minimum number of variables required to "
        "represent the underlying structure of a dataset. In many practical cases, "
        "the ambient dimension is much larger than the true number of degrees of freedom "
        "governing the data."
    )
    add_body_paragraph(
        doc,
        "This report presents the visual results of intrinsic-dimension analysis for "
        "four datasets: a two-dimensional plane embedded in ten dimensions, the Swiss "
        "roll dataset, Gaussian noise, and MNIST handwritten digit images."
    )
    add_body_paragraph(
        doc,
        "Three complementary estimation methods were applied. Principal component analysis "
        "(PCA) was used to examine the number of linear components required to explain "
        "most of the variance, correlation dimension was used to characterize local geometric "
        "scaling, and a k-nearest neighbors (kNN) estimator was used to provide an additional "
        "neighborhood-based estimate of intrinsic dimensionality."
    )

    doc.add_heading("2. Methods Overview", level=1)
    add_bullets(
        doc,
        [
            "PCA was evaluated through cumulative explained variance curves and scree plots.",
            "Correlation dimension was estimated from the slope of the linear scaling region in log-log correlation-integral plots.",
            "A kNN-based Levina-Bickel maximum-likelihood estimator was used to estimate intrinsic dimension from nearest-neighbor distances.",
            "The combined analysis was designed to compare linear variance structure with intrinsic geometric complexity.",
        ],
    )


# ---------------------------------------------------------------------------
# Dataset text content
# ---------------------------------------------------------------------------

_DATASET_TEXTS: Dict[str, Dict[str, List[str]]] = {
    "plane": {
        "pca": [
            "The cumulative explained variance curve shows that almost all variance is captured by the first two principal components. This is exactly the expected result for a linear two-dimensional manifold embedded in a higher-dimensional ambient space.",
            "The figure therefore confirms that PCA successfully recovers the known intrinsic structure of this dataset.",
        ],
        "scree": [
            "The scree plot shows a clear dominance of the first two components, while the remaining components contribute essentially no meaningful variance.",
            "This is the strongest possible visual indication of a low-dimensional linear structure.",
        ],
        "corr": [
            "The correlation-dimension estimate is very close to 2, which matches the theoretical intrinsic dimensionality of the dataset.",
            "The agreement between PCA and correlation dimension validates both the dataset construction and the analysis workflow.",
        ],
        "knn": [
            "The kNN sensitivity plot shows that the estimated intrinsic dimension remains close to 2 over the full tested range of neighborhood sizes.",
            "This provides further support for the conclusion that the dataset is intrinsically two-dimensional and that the estimate is stable.",
        ],
    },
    "swiss_roll": {
        "pca": [
            "The cumulative explained variance plot indicates that multiple principal components are required to preserve most of the variance.",
            "This occurs because the Swiss roll is intrinsically two-dimensional but geometrically curved, and PCA is limited to linear directions.",
        ],
        "scree": [
            "The scree plot does not collapse sharply after the first two components, which reflects the nonlinear geometry of the manifold.",
            "This behavior is expected when a curved low-dimensional structure is analyzed using a linear projection method.",
        ],
        "corr": [
            "The correlation-dimension estimate remains close to 2, which is much more faithful to the true intrinsic structure of the Swiss roll.",
            "This dataset therefore illustrates why geometric methods are especially valuable for nonlinear manifolds.",
        ],
        "knn": [
            "The kNN-based estimates remain close to 2 for all tested values of k, indicating strong stability.",
            "Together with the correlation-dimension result, this confirms that the Swiss roll is intrinsically low-dimensional despite its nonlinear embedding.",
        ],
    },
    "gaussian_noise": {
        "pca": [
            "The cumulative explained variance increases gradually rather than sharply, indicating that variance is distributed across many directions.",
            "This is the expected pattern for isotropic high-dimensional noise without a compact low-dimensional linear subspace.",
        ],
        "scree": [
            "The scree plot shows no pronounced elbow and no small subset of components that dominates the variance structure.",
            "This confirms the absence of a strong linear low-dimensional representation.",
        ],
        "corr": [
            "The correlation-dimension estimate is comparatively high, which is consistent with the noise-like nature of the data.",
            "Together, the PCA and correlation results distinguish Gaussian noise clearly from the more structured datasets.",
        ],
        "knn": [
            "The kNN sensitivity plot remains in a high-dimensional regime across all tested values of k, although the estimates vary more than in the low-dimensional synthetic datasets.",
            "This moderate variability is expected for a dataset that does not lie on a sharply defined low-dimensional manifold.",
        ],
    },
    "mnist": {
        "pca": [
            "The cumulative explained variance curve shows that a very large number of principal components are required to preserve 90%, 95%, and 99% of the total variance.",
            "This reflects the fact that image variability in pixel space is distributed across many linear directions.",
        ],
        "scree": [
            "The scree plot decays gradually, indicating that many components contain non-negligible information.",
            "Such behavior is typical for structured image datasets with rich variability in writing style, thickness, curvature, and placement.",
        ],
        "corr": [
            "The correlation-dimension estimate is much lower than the number of PCA components required for high variance retention.",
            "This gap suggests that MNIST occupies a structured nonlinear region of the ambient space rather than a diffuse cloud filling the full pixel space.",
        ],
        "knn": [
            "The kNN sensitivity analysis produces estimates in the high-teen range and shows good stability across the tested neighborhood sizes.",
            "The kNN estimate lies between the PCA-based and correlation-dimension estimates, which supports the interpretation that MNIST has a structured yet moderately complex intrinsic geometry.",
        ],
    },
}

# (section_title, dataset_title, figure_key_prefix)
_DATASET_ORDER: List[Tuple[str, str, str]] = [
    ("3.1 2D Plane Embedded in 10D", "2D Plane Embedded in 10D", "plane"),
    ("3.2 Swiss Roll", "Swiss Roll", "swiss_roll"),
    ("3.3 Gaussian Noise", "Gaussian Noise", "gaussian_noise"),
    ("3.4 MNIST", "MNIST", "mnist"),
]


def _add_figure_with_text(
    doc: Document,
    numbering: FigureNumbering,
    image_key: str,
    caption: str,
    body_paragraphs: List[str],
) -> None:
    """
    Insert a figure with its caption, followed by body paragraphs.
    """
    add_figure(
        doc,
        numbering,
        FIGURES_DIR / FIGURE_FILES[image_key],
        caption,
    )
    for paragraph_text in body_paragraphs:
        add_body_paragraph(doc, paragraph_text)


def add_dataset_results(
    doc: Document,
    numbering: FigureNumbering,
    section_title: str,
    dataset_title: str,
    pca_cumulative_key: str,
    scree_key: str,
    correlation_key: str,
    knn_key: str,
    pca_text: List[str],
    scree_text: List[str],
    corr_text: List[str],
    knn_text: List[str],
) -> None:
    """
    Add all figure-based results for a single dataset.
    """
    doc.add_heading(section_title, level=2)

    add_body_paragraph(
        doc,
        f"This subsection summarizes the visual results obtained for the {dataset_title} dataset."
    )

    _add_figure_with_text(
        doc, numbering, pca_cumulative_key,
        f"Cumulative explained variance curve for the {dataset_title} dataset.",
        pca_text,
    )
    _add_figure_with_text(
        doc, numbering, scree_key,
        f"Scree plot for the {dataset_title} dataset.",
        scree_text,
    )
    _add_figure_with_text(
        doc, numbering, correlation_key,
        f"Correlation-dimension log-log plot for the {dataset_title} dataset.",
        corr_text,
    )
    _add_figure_with_text(
        doc, numbering, knn_key,
        f"kNN sensitivity analysis for the {dataset_title} dataset across multiple values of k.",
        knn_text,
    )


def add_knn_section(doc: Document) -> None:
    """
    Add the report-ready kNN methodology and interpretation section.
    """
    doc.add_heading("4. k-Nearest Neighbors Based Intrinsic Dimension Estimation", level=1)

    add_body_paragraph(
        doc,
        "To complement the PCA-based analysis and the correlation-dimension approach, "
        "a third estimator based on k-nearest neighbors was applied. This method provides "
        "an additional local geometric perspective on intrinsic dimensionality and does not "
        "rely on global linear projections."
    )
    add_body_paragraph(
        doc,
        "In this project, intrinsic dimension was estimated using the Levina-Bickel "
        "maximum-likelihood estimator. For each point, the estimator uses the distances "
        "to its nearest neighbors and computes a local estimate of intrinsic dimension. "
        "These local estimates are then aggregated across the dataset to obtain a global estimate."
    )
    add_body_paragraph(
        doc,
        "The same preprocessing pipeline used in the rest of the project was applied before "
        "the kNN analysis. Constant features were removed, the data were standardized, and "
        "no additional normalization was applied. For MNIST, the estimator was computed on "
        "a random sample of 1000 observations in order to maintain computational efficiency."
    )
    add_body_paragraph(
        doc,
        "A headline estimate was computed with k=10, and a sensitivity analysis was performed "
        "over the range k in {5, 8, 10, 12, 15, 20}. This made it possible to evaluate the "
        "stability of the estimator with respect to neighborhood size."
    )

    doc.add_heading("4.1 kNN Summary Table", level=2)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = [
        "Dataset",
        "Samples Used",
        "Processed Dimension",
        "kNN Estimate (k=10)",
        "Mean Across k",
        "Std. Across k",
    ]

    header_cells = table.rows[0].cells
    for index, text in enumerate(headers):
        cell = header_cells[index]
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    rows = [
        ("MNIST (sampled)", "1000", "599", "16.71", "17.23", "0.83"),
        ("Gaussian Noise", "2000", "20", "17.68", "17.89", "1.73"),
        ("2D Plane Embedded in 10D", "2000", "10", "2.22", "2.25", "0.18"),
        ("Swiss Roll", "2000", "3", "2.19", "2.24", "0.19"),
    ]

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                format_paragraph(paragraph, space_after=2, line_spacing=1.05)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

    doc.add_paragraph()

    add_body_paragraph(
        doc,
        "The kNN-based results are highly consistent with the overall interpretation of the datasets. "
        "For the 2D plane embedded in 10D and for the Swiss roll dataset, the kNN estimates were both "
        "very close to 2, which strongly supports the conclusion that both datasets are intrinsically "
        "two-dimensional despite differences in their global geometry."
    )
    add_body_paragraph(
        doc,
        "For Gaussian noise, the kNN estimate remained high, which is fully consistent with the absence "
        "of a compact low-dimensional manifold. For MNIST, the estimate was much lower than the ambient "
        "dimension but substantially larger than for the simple synthetic manifolds, indicating a structured "
        "yet nontrivial intrinsic geometry."
    )


def add_global_comparison(doc: Document, numbering: FigureNumbering) -> None:
    """
    Add cross-dataset comparison figures.
    """
    doc.add_heading("5. Global Comparison Across Datasets", level=1)

    add_body_paragraph(
        doc,
        "After analyzing each dataset individually, it is useful to compare the final "
        "PCA-based, correlation-dimension, and kNN-based estimates across all datasets."
    )

    add_figure(
        doc,
        numbering,
        FIGURES_DIR / FIGURE_FILES["comparison_all_methods"],
        "Comparison of PCA, correlation-dimension, and kNN-based dimensionality estimates across all datasets.",
        width=6.4,
    )

    add_body_paragraph(
        doc,
        "The full comparison clearly shows that the three methods capture different aspects of structure. "
        "PCA is variance-oriented and linear, correlation dimension is locally geometric, and the kNN estimator "
        "provides a neighborhood-based estimate that often lies between the other two."
    )

    add_figure(
        doc,
        numbering,
        FIGURES_DIR / FIGURE_FILES["comparison_synthetic_all_methods"],
        "Comparison of PCA, correlation-dimension, and kNN estimates restricted to the synthetic datasets.",
        width=6.4,
    )

    add_body_paragraph(
        doc,
        "The synthetic-only comparison makes the relationships among the controlled datasets easier to interpret. "
        "It shows especially clearly that the 2D plane and the Swiss roll are intrinsically close to two-dimensional, "
        "while Gaussian noise remains high-dimensional."
    )


def add_mnist_visual_reference(doc: Document, numbering: FigureNumbering) -> None:
    """
    Add the MNIST sample figure.
    """
    doc.add_heading("6. Visual Reference for MNIST", level=1)

    add_body_paragraph(
        doc,
        "A sample of MNIST images is included below to provide visual context for the "
        "type of data analyzed in the experiment."
    )

    add_figure(
        doc,
        numbering,
        FIGURES_DIR / FIGURE_FILES["mnist_samples"],
        "Random sample of MNIST handwritten digit images.",
        width=6.4,
    )

    add_body_paragraph(
        doc,
        "Although each image is represented numerically as a 784-dimensional vector, "
        "the visual variation is governed by a much smaller set of semantic factors, "
        "including digit identity, stroke thickness, slant, and writing style."
    )


def add_conclusion(doc: Document) -> None:
    """
    Add the final discussion and conclusion.
    """
    doc.add_heading("7. Discussion and Conclusion", level=1)

    add_body_paragraph(
        doc,
        "The results demonstrate that intrinsic dimension is not a single quantity that "
        "all methods recover identically. Instead, the estimate depends on what aspect "
        "of structure is being measured."
    )

    add_bullets(
        doc,
        [
            "PCA is most informative when the data are approximately linear and the goal is variance preservation.",
            "Correlation dimension is better suited to capturing nonlinear geometric structure.",
            "The kNN estimator provides an additional neighborhood-based perspective and complements both PCA and correlation dimension.",
            "Agreement between the methods is strongest for simple linear datasets, while disagreement becomes informative for nonlinear and complex real-world data.",
        ],
    )

    add_body_paragraph(
        doc,
        "Overall, the 2D plane embedded in 10D served as a clear sanity check, the Swiss roll highlighted "
        "the limitations of linear methods, Gaussian noise behaved as a high-dimensional unstructured dataset, "
        "and MNIST revealed a structured but nontrivial intrinsic geometry."
    )

    add_body_paragraph(
        doc,
        "Taken together, the three estimators provide a richer and more balanced view of dataset complexity "
        "than any single method alone."
    )


def build_document() -> Document:
    """
    Build the full report document.
    """
    doc = Document()
    configure_page_layout(doc)
    apply_base_styles(doc)
    add_header_and_footer(doc)

    numbering = FigureNumbering()

    add_title_page(doc)
    add_intro_and_methods(doc)

    doc.add_heading("3. Results and Figure Analysis", level=1)

    for section_title, dataset_title, key in _DATASET_ORDER:
        texts = _DATASET_TEXTS[key]
        add_dataset_results(
            doc=doc,
            numbering=numbering,
            section_title=section_title,
            dataset_title=dataset_title,
            pca_cumulative_key=f"pca_cumulative_{key}",
            scree_key=f"scree_{key}",
            correlation_key=f"correlation_{key}",
            knn_key=f"knn_{key}",
            pca_text=texts["pca"],
            scree_text=texts["scree"],
            corr_text=texts["corr"],
            knn_text=texts["knn"],
        )

    add_knn_section(doc)
    add_global_comparison(doc, numbering)
    add_mnist_visual_reference(doc, numbering)
    add_conclusion(doc)

    return doc


def main() -> None:
    """
    Validate figure availability and create the DOCX file.
    """
    missing_files: List[str] = []

    for file_name in FIGURE_FILES.values():
        path = FIGURES_DIR / file_name
        if not path.exists():
            missing_files.append(str(path))

    if missing_files:
        print("The following figure files are missing:")
        for item in missing_files:
            print(f" - {item}")
        return

    document = build_document()
    document.save(str(OUTPUT_DOCX))

    print("DOCX report created successfully.")
    print(f"Output file: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
